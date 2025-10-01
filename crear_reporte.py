import pandas as pd
from docx import Document
from sklearn.linear_model import LinearRegression
import numpy as np
import sys

# --- INICIO DE FUNCIONES AUXILIARES ---

def generar_proyeccion(df_datos, columna_consumo):
    """
    Calcula la proyección a 6 y 12 meses para una columna de consumo.
    Devuelve las predicciones o None si no hay suficientes datos.
    """
    consumo_mensual = df_datos.groupby('Fecha')[columna_consumo].sum().reset_index()
    
    if len(consumo_mensual) < 2:
        return None

    consumo_mensual['Mes_Num'] = np.arange(len(consumo_mensual))
    modelo = LinearRegression()
    modelo.fit(consumo_mensual[['Mes_Num']], consumo_mensual[columna_consumo])
    
    ultimo_mes_num = consumo_mensual['Mes_Num'].max()
    meses_futuros = np.array([[ultimo_mes_num + 6], [ultimo_mes_num + 12]])
    predicciones = modelo.predict(meses_futuros).round(0)
    
    if predicciones[0] >= 0 and predicciones[1] >= 0:
        return predicciones
    return None

def agregar_tabla_a_word(documento, df_datos, titulo=""):
    """
    Añade un título y un DataFrame de pandas como una tabla a un documento de Word.
    """
    if titulo:
        documento.add_paragraph(titulo, style='Heading 4')
        
    if df_datos.empty:
        documento.add_paragraph("No se generaron datos para esta tabla.")
        return

    tabla = documento.add_table(rows=1, cols=len(df_datos.columns))
    tabla.style = 'Table Grid'
    
    for i, nombre_columna in enumerate(df_datos.columns):
        tabla.cell(0, i).text = str(nombre_columna)
        
    for _, fila in df_datos.iterrows():
        celdas_fila = tabla.add_row().cells
        for i, valor in enumerate(fila):
            if isinstance(valor, (int, float, np.integer)):
                celdas_fila[i].text = f"{int(valor):,}".replace(',', '.')
            else:
                celdas_fila[i].text = str(valor)

# --- INICIO DEL SCRIPT PRINCIPAL ---

# --- Configuración ---
archivo_entrada = 'resumen_consolidado.xlsx'
archivo_salida_word = 'Reporte_Analisis_Consumo_Anual_Completo.docx'

# --- 1. Cargar y Preparar los Datos ---
try:
    print(f"Leyendo la hoja 'Resumen Detallado' del archivo: {archivo_entrada}...")
    df = pd.read_excel(archivo_entrada, sheet_name='Resumen Detallado')
except FileNotFoundError:
    print(f"❌ ERROR: No se encontró el archivo '{archivo_entrada}'.")
    sys.exit()
except Exception as e:
    print(f"❌ Ocurrió un error inesperado al leer el archivo: {e}")
    sys.exit()

print("Convirtiendo periodos a fechas...")
meses_es = {
    'Enero': 'January', 'Febrero': 'February', 'Marzo': 'March', 'Abril': 'April',
    'Mayo': 'May', 'Junio': 'June', 'Julio': 'July', 'Agosto': 'August',
    'Septiembre': 'September', 'Octubre': 'October', 'Noviembre': 'November', 'Diciembre': 'December'
}
periodo_traducido = df['Periodo'].str.split().str[0].map(meses_es)
año_str = df['Periodo'].str.split().str[1]
periodo_en_ingles = periodo_traducido + ' ' + año_str
df['Fecha'] = pd.to_datetime(periodo_en_ingles)
df['Año'] = df['Fecha'].dt.year

# --- 2. Preparar el Documento de Word ---
documento = Document()
años_en_datos = sorted(df['Año'].unique())
if len(años_en_datos) == 1:
    titulo_principal = f'Análisis de Consumo y Proyecciones {años_en_datos[0]}'
else:
    titulo_principal = f'Análisis de Consumo y Proyecciones {años_en_datos[0]}-{años_en_datos[-1]}'
documento.add_heading(titulo_principal, level=1)

columnas_analisis = ['Total Imp. B/N', 'Total Copias B/N', 'Total Imp. Color', 'Total Copias Color']

# --- 3. Bucle para Analizar Cada Año por Separado ---
for año in años_en_datos:
    print(f"\n--- Procesando datos para el año: {año} ---")
    documento.add_heading(f'Resultados para el Año {año}', level=2)
    
    df_año = df[df['Año'] == año].copy()

    # --- NUEVA LÓGICA PARA CÁLCULO CONDICIONAL ---
    print(f"Calculando sumas y promedios condicionales para {año}...")

    # 1. Agrupar para obtener la suma total Y el conteo de meses únicos por agencia
    analisis_agencia = df_año.groupby('Agencia').agg(
        **{col: pd.NamedAgg(column=col, aggfunc='sum') for col in columnas_analisis},
        Meses_Registrados=pd.NamedAgg(column='Fecha', aggfunc='nunique')
    )

    # 2. Crear una columna que indique el tipo de cálculo
    analisis_agencia['Tipo de Cifra'] = np.where(
        analisis_agencia['Meses_Registrados'] == 12,
        'Promedio Mensual', # Si hay 12 meses
        'Suma Parcial'      # Si hay menos de 12 meses
    )

    # 3. Aplicar la división solo a las agencias con 12 meses
    for col in columnas_analisis:
        analisis_agencia[col] = np.where(
            analisis_agencia['Meses_Registrados'] == 12,
            (analisis_agencia[col] / 12), # Calcula el promedio
            analisis_agencia[col]         # Mantiene la suma total
        ).round(0) # Redondear el resultado final

    # 4. Limpiar el DataFrame antes de guardarlo en Word
    analisis_agencia.drop(columns=['Meses_Registrados'], inplace=True)
    analisis_agencia.reset_index(inplace=True)
    
    # Mover la columna 'Tipo de Cifra' para que aparezca después de 'Agencia'
    columnas_ordenadas = ['Agencia', 'Tipo de Cifra'] + columnas_analisis
    analisis_agencia = analisis_agencia[columnas_ordenadas]
    
    # --- FIN DE LA NUEVA LÓGICA ---

    parrafo_descripcion = (
        f"La siguiente tabla muestra el consumo para {año}. Las cifras representan el 'Promedio Mensual' "
        "para agencias con 12 meses de datos, o la 'Suma Parcial' para agencias con datos incompletos."
    )
    documento.add_paragraph(parrafo_descripcion)
    agregar_tabla_a_word(documento, analisis_agencia)


# --- 4. Proyecciones Globales (utilizando todos los datos) ---
print("\n--- Realizando proyecciones globales con todos los datos ---")
documento.add_heading('Proyecciones Globales a Futuro', level=2)

# a) Proyección por Tipo de Consumo (General)
datos_proy_tipo = []
for columna in columnas_analisis:
    predicciones = generar_proyeccion(df, columna)
    if predicciones is not None:
        datos_proy_tipo.append({
            'Tipo de Consumo': columna,
            'Proyección a 6 Meses': predicciones[0],
            'Proyección a 1 Año': predicciones[1]
        })
df_proyecciones_tipo = pd.DataFrame(datos_proy_tipo)
agregar_tabla_a_word(documento, df_proyecciones_tipo, "Proyección general por tipo de consumo")

# b) Proyección por Agencia y por Tipo
datos_proy_agencia = []
for agencia, df_agencia in df.groupby('Agencia'):
    for columna in columnas_analisis:
        predicciones = generar_proyeccion(df_agencia, columna)
        if predicciones is not None:
            datos_proy_agencia.append({
                'Agencia': agencia,
                'Tipo de Consumo': columna,
                'Proyección a 6 Meses': predicciones[0],
                'Proyección a 1 Año': predicciones[1]
            })
df_proyecciones_agencia = pd.DataFrame(datos_proy_agencia)
agregar_tabla_a_word(documento, df_proyecciones_agencia, "Proyección detallada por agencia y tipo de consumo")


# --- 5. Guardar el Documento ---
try:
    documento.save(archivo_salida_word)
    print("-" * 30)
    print(f"✅ ¡Reporte anual generado exitosamente en '{archivo_salida_word}'!")
except Exception as e:
    print(f"❌ ERROR: No se pudo guardar el archivo de Word. Asegúrate de que no esté abierto.")
    print(f"   Detalle del error: {e}")